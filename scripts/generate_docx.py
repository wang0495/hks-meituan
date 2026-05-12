from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 设置中文字体
doc.styles['Normal'].font.name = 'Microsoft YaHei'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# 标题
title = doc.add_heading('CityFlow 分工与成本控制指南', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 副标题
subtitle = doc.add_paragraph('版本: v1.0 | 日期: 2024-05-12')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(10)
subtitle.runs[0].font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# 一、关键信息
doc.add_heading('一、关键信息', 1)
doc.add_heading('1. 当前状态', 2)
doc.add_paragraph('正确率: 63%（19/30场景通过）')
doc.add_paragraph('架构: 3个Agent（IntentAgent → FeasibilityAgent → Solver）')
doc.add_paragraph('问题: 11个失败场景，主要是该拒绝时没拒绝')

doc.add_heading('2. 为什么现在分工？', 2)
table = doc.add_table(rows=3, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '阶段'
hdr_cells[1].text = '做什么'
hdr_cells[2].text = '为什么'
row_cells = table.rows[1].cells
row_cells[0].text = '前两周'
row_cells[1].text = '快速搭框架'
row_cells[2].text = '验证Agent架构是否可行'
row_cells = table.rows[2].cells
row_cells[0].text = '现在'
row_cells[1].text = '精细化打磨'
row_cells[2].text = '框架OK，需要规模化'

doc.add_heading('3. 核心问题', 2)
doc.add_paragraph('FeasibilityAgent漏检: 4个场景该拒绝没拒绝', style='List Bullet')
doc.add_paragraph('营业时间过滤: 深夜场景返回已关门的店', style='List Bullet')
doc.add_paragraph('地理连续性: 美食路线跑到郊区30km', style='List Bullet')

# 二、成本控制
doc.add_heading('二、成本控制方案', 1)
doc.add_heading('不推荐（太贵）', 2)
doc.add_paragraph('Claude Code: 20美元/月 + API费', style='List Bullet')
doc.add_paragraph('GitHub Copilot: 10美元/月', style='List Bullet')
doc.add_paragraph('豆包: 禁止使用', style='List Bullet')

doc.add_heading('推荐（省钱）', 2)
p = doc.add_paragraph()
p.add_run('方案A: Coding Plan（最低成本）').bold = True
doc.add_paragraph('工具: Cursor免费版 + DeepSeek V4 Flash')
doc.add_paragraph('成本: 几乎免费（1元/百万token）')

p = doc.add_paragraph()
p.add_run('方案B: ccswitch切换模型').bold = True
doc.add_paragraph('原理: 平时用DeepSeek，复杂问题才用Claude')
doc.add_paragraph('成本: 节省80%以上')

# 三、同学A任务
doc.add_heading('三、同学A任务（会一点代码）', 1)
doc.add_heading('任务1: 修复3个bug（3天）', 2)
doc.add_paragraph('backend/agents/__init__.py - 修改FeasibilityAgent')
doc.add_paragraph('backend/services/solver.py - 限制10km搜索半径')
doc.add_paragraph('backend/services/filters.py - 深夜24h过滤')

doc.add_heading('任务2: 写测试框架（2天）', 2)
doc.add_paragraph('test_layer1.py - 单元测试')
doc.add_paragraph('test_layer2.py - Agent协作测试')
doc.add_paragraph('test_layer3.py - 端到端测试')

doc.add_heading('任务3: 自动生成场景（1天）', 2)
doc.add_paragraph('生成100个边界场景（预算/时间边界）')

# 四、同学B任务
doc.add_heading('四、同学B任务（完全不会代码）', 1)
doc.add_heading('任务1: 人工审核失败案例（1天）', 2)
doc.add_paragraph('看11个失败场景，判断LLM评分是否找茬')

doc.add_heading('任务2: 写真实场景（2天）', 2)
doc.add_paragraph('写20个"人话版"场景（名称+输入+预期）')

doc.add_heading('任务3: 竞品对比（1天）', 2)
doc.add_paragraph('测美团/高德/小红书，记录差异')

# 五、实时维护架构图
doc.add_heading('五、实时维护架构图', 1)
doc.add_paragraph('方法: 每次commit后自动更新架构图')
doc.add_paragraph('文件: docs/architecture.md + docs/architecture.png')
doc.add_paragraph('工具: git hook + mermaid自动生成')

# 六、每周同步
doc.add_heading('六、每周同步模板', 1)
doc.add_heading('同学A汇报（数据+成本）', 2)
doc.add_paragraph('修了X个bug，正确率 63% → 71%')
doc.add_paragraph('API花费: ￥X元（控制预算内）')

doc.add_heading('同学B汇报（体验）', 2)
doc.add_paragraph('审了X个案例，Y个真有问题')
doc.add_paragraph('竞品对比发现: 美团在XX方面更好')

# 七、工具链
doc.add_heading('七、工具链推荐', 1)
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '角色'
hdr_cells[1].text = '推荐工具'
hdr_cells[2].text = '成本'
data = [
    ['同学A', 'Cursor免费版 + DeepSeek', '免费/￥1每百万token'],
    ['同学A', 'ccswitch', '免费'],
    ['同学B', '不需要编程工具', '-'],
    ['共同', 'LongCat（公测免费）', '免费']
]
for i, row_data in enumerate(data, 1):
    row_cells = table.rows[i].cells
    for j, text in enumerate(row_data):
        row_cells[j].text = text

# 八、避坑
doc.add_heading('八、避坑提醒', 1)
doc.add_heading('成本控制红线', 2)
doc.add_paragraph('不要用Claude Code月租版', style='List Bullet')
doc.add_paragraph('不要用GitHub Copilot', style='List Bullet')
doc.add_paragraph('不要在Cursor里开Pro', style='List Bullet')

doc.add_heading('协作红线', 2)
doc.add_paragraph('每天同步一次', style='List Bullet')
doc.add_paragraph('不懂就问，别猜', style='List Bullet')

# 九、快速启动
doc.add_heading('九、快速启动清单', 1)
doc.add_heading('Day 1', 2)
doc.add_paragraph('同学A: 安装Cursor，配置DeepSeek API', style='List Bullet')
doc.add_paragraph('同学B: 看失败案例分析', style='List Bullet')
doc.add_paragraph('一起: 确认分工', style='List Bullet')

doc.add_heading('Day 2-3', 2)
doc.add_paragraph('同学A: 修FeasibilityAgent', style='List Bullet')
doc.add_paragraph('同学B: 完成11个案例审核', style='List Bullet')

doc.add_heading('Day 4-5', 2)
doc.add_paragraph('同学A: 写测试框架', style='List Bullet')
doc.add_paragraph('同学B: 写20个真实场景', style='List Bullet')

# 保存
output_path = 'C:\\Users\\wang\\Desktop\\CityFlow_分工与成本控制指南.docx'
doc.save(output_path)
print(f'文件已保存: {output_path}')
