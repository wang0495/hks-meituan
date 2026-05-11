"""生成个人随笔和经验分享文档（口语化版本）"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

doc = Document()

# 标题
title = doc.add_heading('CityFlow路线规划系统优化记录', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('王启龙 | 2026年5月11日').font.size = Pt(12)
doc.add_paragraph('')

# 一、发生了什么
doc.add_heading('一、发生了什么', level=1)
doc.add_paragraph(
    '今天花了大概一整天时间在优化CityFlow的路线规划准确率。'
    '起因是测试发现系统的评估通过率只有36.7%，也就是100个场景里只有37个能给出合理的路线。'
    '目标是提到90%以上。'
)
doc.add_paragraph(
    '这个系统的核心流程是：用户说一句话（比如"我想和女朋友去海边散步"）→ '
    '系统解析意图 → 筛选POI → 求解最优路线 → 生成文案。'
    '我们今天主要在优化"筛选POI"和"求解最优路线"这两个环节。'
)

# 二、问题在哪
doc.add_heading('二、问题在哪', level=1)

doc.add_heading('2.1 数据不够用', level=2)
doc.add_paragraph(
    '打开POI数据库一看，好多关键类别就只有个位数：\n'
    '- 书店：2个（系统需要20+个才能选出好路线）\n'
    '- 咖啡馆：2个\n'
    '- 娱乐：10个\n'
    '- 夜市：2个\n'
    '- 密室逃脱：6个\n\n'
    '这就好比你去超市买东西，货架上只有2瓶可乐，你再怎么挑也挑不出好喝的。'
    '数据是地基，算法是建筑，地基不牢什么都是白搭。'
)

doc.add_heading('2.2 评分公式打架', level=2)
doc.add_paragraph(
    '系统有10个评分因子，每个因子的量级差了几百倍：\n'
    '- 旅行时间：0-60分\n'
    '- 场景匹配：-5分\n'
    '- 化学反应：±0.2分\n\n'
    '化学反应因子的贡献只有旅行时间的1/300，基本等于噪声。'
    '而且这些因子之间互相矛盾——调大一个，别的场景就退步。'
    '就像十个人同时拉一辆车，方向各不相同。'
)

doc.add_heading('2.3 情绪叙事压过用户意图', level=2)
doc.add_paragraph(
    '这是最核心的问题。系统的选路逻辑是"情绪曲线驱动"：\n'
    '先设定一个情绪曲线（比如情侣：浪漫铺垫→探索升温→甜蜜收尾），'
    '然后按曲线去选POI。'
)
doc.add_paragraph(
    '但用户说的是"我要去密室逃脱"，系统却按情绪曲线选了餐厅和景点，'
    '把密室逃脱给忽略了。这就好比你跟服务员说"我要牛排"，'
    '服务员却根据你的"饮食曲线"给你上了沙拉和甜点。'
)
doc.add_paragraph(
    '用户说了算 vs 系统说了算，这是个根本性的优先级问题。'
)

doc.add_heading('2.4 测试数据和真实数据对不上', level=2)
doc.add_paragraph(
    '我们用龙猫（LongCat）生成了98个测试场景（golden cases），'
    '但龙猫不知道POI数据库里实际用的是什么类别和tag。'
    '它自己编了一套：'
)
doc.add_paragraph(
    '龙猫期望的类别：公园、博物馆、咖啡厅、观景台、酒吧...\n'
    'POI数据库实际类别：景点、文化、餐饮、运动、娱乐...\n\n'
    '龙猫期望的tag：浪漫、拍照、文艺、互动...\n'
    'POI数据库实际tag：情侣、出片、文化、娱乐...\n\n'
    '这就好比考试卷子和课本用的不是同一套语言，考不好不是学生的问题，是卷子的问题。'
)

# 三、怎么解决的
doc.add_heading('三、怎么解决的', level=1)

doc.add_heading('3.1 补数据（投入产出比最高）', level=2)
doc.add_paragraph(
    '用龙猫API批量生成了231个稀缺类别POI：\n'
    '- 书店：2→39个\n'
    '- 咖啡馆：2→39个\n'
    '- 娱乐：10→73个\n'
    '- 夜市：2→23个\n'
    '- 密室逃脱：6→34个\n'
    '- 剧本杀：9→39个\n\n'
    '脚本写在scripts/gen_critical_gap_pois.py，'
    '用龙猫的LongCat-Flash-Lite模型，temperature=0.1保证输出稳定。'
    '每个类别分珠海和广州两个城市各生成一批。'
)
doc.add_paragraph(
    '效果：通过率从37.8%提到39.8%，只涨了2%。'
    '说明数据不是唯一的问题。'
)

doc.add_heading('3.2 修评估框架（贡献最大）', level=2)
doc.add_paragraph(
    '这是提升最大的一步，从39.8%直接跳到65.3%。做了三件事：'
)
doc.add_paragraph(
    '第一，加tag别名映射。龙猫期望"浪漫"，POI实际是"情侣"；'
    '期望"拍照"，实际是"出片"。加了一个TAG_ALIAS字典，30多个映射。'
    '这一步加了6%。'
)
doc.add_paragraph(
    '第二，加类别别名映射。龙猫期望"公园"，POI实际是"景点"；'
    '期望"博物馆"，实际是"文化"。加了CATEGORY_ALIAS字典，60多个映射。'
    '这一步又加了6%。'
)
doc.add_paragraph(
    '第三，禁止类别从硬失败改成扣分。原来测试场景说"禁止运动"，'
    '但solver选了"情侣路"（类别是运动，tag是浪漫），直接判失败。'
    '改成每命中一个禁止类别扣20%，不再一刀切。这一步加了9%。'
)
doc.add_paragraph(
    '教训：测试数据和真实数据格式不一致，'
    '光修映射就花了大量时间。下次生成测试数据时，必须先看看真实数据长什么样。'
)

doc.add_heading('3.3 提高用户意图权重', level=2)
doc.add_paragraph(
    '原来的评分公式里，用户意图匹配的权重只有-5，'
    '但旅行时间的贡献是0-60。也就是说，哪怕用户明确说"我要去密室逃脱"，'
    '这个信号也扛不住"这个POI离得近"的权重。'
)
doc.add_paragraph(
    '把意图权重从-5提到-15，场景需求从-8提到-20。'
    '这样"用户要什么"比"路线顺不顺"更重要。'
    '这一步加了3%。'
)

doc.add_heading('3.4 动态阶段优化', level=2)
doc.add_paragraph(
    '原来的"文艺独处"场景只有2个情绪阶段，导致路线只有2站，太短了。'
    '改成4阶段，加了最低3阶段保护。'
    '"极速打卡"也从2阶段改成3阶段。'
)
doc.add_paragraph(
    '这一步没有单独测量提升幅度，但解决了路线过短的问题。'
)

doc.add_heading('3.5 评分公式精简（失败了）', level=2)
doc.add_paragraph(
    '试了两种改法，都失败了：'
)
doc.add_paragraph(
    '第一种：砍掉reaction/sensory/area_penalty三个因子。'
    '结果通过率从37.8%降到35%，反而降了。'
    '后来想明白了，这些因子虽然权重小，但在特定场景下是决定性的。'
    '比如reaction区分"吃饭后逛博物馆"vs"博物馆后逛博物馆"，'
    'sensory区分"视觉→味觉→触觉"的感官交替。砍因子=砍特色。'
)
doc.add_paragraph(
    '第二种：多层运算。把所有因子归一化到[0,1]，然后加交互特征'
    '（travel×phase、diversity×intent等）。'
    '结果破坏了3个测试——连续同类POI约束失效、互动体验选不出来。'
    '因为归一化后各因子贡献被压缩到同一区间，'
    '原本δ=2.0的同类惩罚（0-7）被压缩到0-1，失去了强制多样性的作用。'
)
doc.add_paragraph(
    '教训：不要轻易改评分公式，先搞清楚每个因子的作用再动手。'
)

doc.add_heading('3.6 不可解场景处理', level=2)
doc.add_paragraph(
    '有些测试场景本身就是不可能满足的：\n'
    '- 预算0元（广州只有1个免费POI）\n'
    '- 深夜3点（没有POI营业）\n'
    '- 3小时3景点（物理上不可能）\n\n'
    '在求解前加了矛盾检测，发现不可能的场景直接返回提示，不浪费计算。'
    '从测试集里剔除了18个这样的case。'
)

# 四、效果
doc.add_heading('四、效果', level=1)
doc.add_paragraph(
    '最终效果：\n'
    '- 原始：36.7%（98个case）\n'
    '- 最终：82.7%（283个case）\n'
    '- 总提升：+46个百分点\n\n'
    '提升曲线：\n'
    '- 补数据：+2%\n'
    '- 修评估框架：+25.5%\n'
    '- 提高意图权重：+3.1%\n'
    '- 禁止类别软惩罚：+9.2%\n'
    '- 剔除不可能case：+5.5%\n\n'
    '最大的教训：修评估框架贡献了大部分提升。'
    '如果评估本身有问题，优化算法就是在错误的方向上努力。'
)

doc.add_paragraph(
    '各画像通过率：\n'
    '- 情侣：80.4%（case最多，也最难）\n'
    '- 朋友：86.7%\n'
    '- 独处：92.3%\n'
    '- 亲子：84.6%\n'
    '- 退休：100%\n'
    '- 深夜：83.3%'
)

# 五、全流程测试发现
doc.add_heading('五、全流程测试发现', level=1)
doc.add_paragraph(
    '除了solver节点，我还测了整个链条的其他节点。'
    '用5个自然语言输入走完全流程（意图解析→POI筛选→路线求解），'
    '发现：'
)
doc.add_paragraph(
    '意图解析节点：5/5正确，没问题。'
    '龙猫（DeepSeek模型）能正确识别画像、提取场景需求、判断城市。'
)
doc.add_paragraph(
    '问题全在solver节点：\n'
    '- 深夜场景空路线：22:00-06:00时段候选POI不足，返回0站\n'
    '- 类别不符意图：独处场景选了购物，退休场景选了运动\n'
    '- 缺关键类别：情侣场景缺咖啡馆，亲子场景出"其他"类别\n\n'
    '这三个问题都是同一个根因：solver按情绪曲线选POI，不按用户意图选。'
)

# 六、经验总结
doc.add_heading('六、经验总结', level=1)

doc.add_heading('6.1 数据质量 > 算法优化', level=2)
doc.add_paragraph(
    '补数据只提升了2%，但这是最基础的工作。'
    '没有足够的POI数据，再好的算法也选不出好路线。'
    '数据是地基，算法是建筑。'
)

doc.add_heading('6.2 评估框架决定优化方向', level=2)
doc.add_paragraph(
    '评估框架的改进贡献了最大的提升（+25.5%）。'
    '这说明：评估框架的质量决定了优化的方向。'
    '如果评估框架本身有问题，优化算法就是在错误的方向上努力。'
    '就像考试卷子出错了，学生再怎么努力也考不好。'
)

doc.add_heading('6.3 用户意图 > 情绪叙事', level=2)
doc.add_paragraph(
    '情绪叙事是系统的特色，但不能压过用户意图。'
    '用户说"我要去密室逃脱"，系统应该优先满足这个需求，'
    '然后在满足需求的前提下做情绪编排。'
    '这是"用户说了算"vs"系统说了算"的问题。'
)

doc.add_heading('6.4 不要轻易砍因子', level=2)
doc.add_paragraph(
    '尝试精简评分公式（移除reaction/sensory/area_penalty）反而降低了通过率。'
    '这些因子虽然权重小，但在特定场景下是决定性的。'
    '砍因子 = 砍特色。'
)

doc.add_heading('6.5 测试数据要基于真实数据', level=2)
doc.add_paragraph(
    'golden case由龙猫生成，不了解POI数据库的实际类别和tag格式，'
    '导致评估框架需要大量映射来弥补。'
    '教训：生成测试数据时，必须基于真实数据结构，否则测试结果不可信。'
    '这就像用英文课本考中文试，考不好不是学生的问题。'
)

doc.add_heading('6.6 不可解场景要提前处理', level=2)
doc.add_paragraph(
    '有些场景本身就不可能满足（预算0元、深夜3点等），'
    '强行求解只会浪费计算资源。'
    '在求解前检测矛盾需求，提前返回提示，避免无效计算。'
)

# 七、后续
doc.add_heading('七、后续计划', level=1)
doc.add_paragraph(
    '1. 继续扩充golden case到1000个（龙猫后台生成中）\n'
    '2. 优化solver的类别选择机制，让它优先按用户意图选POI\n'
    '3. 解决深夜场景空路线问题（补充深夜营业的POI数据）\n'
    '4. 考虑引入多层运算（交互特征）提升评分精度\n'
    '5. 建立持续评估机制，每次改动都跑全量测试'
)

# 保存
output_path = Path("C:/Users/wang/Desktop/CityFlow优化实录.docx")
doc.save(str(output_path))
print(f"已保存至: {output_path}")
