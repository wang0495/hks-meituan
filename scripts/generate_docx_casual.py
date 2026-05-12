from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 设置中文字体
doc.styles['Normal'].font.name = 'Microsoft YaHei'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# 标题
title = doc.add_heading('CityFlow 干活手册', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('别花冤枉钱版 | 2024.05.12')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(10)
subtitle.runs[0].font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# 一、现在啥情况
doc.add_heading('一、现在啥情况', 1)

p = doc.add_paragraph()
p.add_run('正确率：').bold = True
p.add_run('63%（30个场景过了19个）')

p = doc.add_paragraph()
p.add_run('啥架构：').bold = True
p.add_run('3个Agent串联 - IntentAgent先筛一遍，FeasibilityAgent再看能不能做，最后Solver排路线')

p = doc.add_paragraph()
p.add_run('主要问题：').bold = True
doc.add_paragraph('有11个场景挂了，基本都是该拒绝的时候没拒绝，硬给塞了个路线', style='List Bullet')
doc.add_paragraph('比如用户说"我想蹦迪"，城里根本没有酒吧，系统应该直接说"做不了"，结果给推荐了几个公园...', style='List Bullet')

doc.add_heading('为啥现在才找人分工？', 2)
doc.add_paragraph('前两周：我一个人疯狂搭框架，验证这条路能不能走通')
doc.add_paragraph('现在：框架跑通了，正确率有63%，但需要规模化打磨，一个人搞不过来')

doc.add_page_break()

# 二、怎么省钱
doc.add_heading('二、怎么省钱（重点！）', 1)

doc.add_heading('别用的（贵且没必要）', 2)
doc.add_paragraph('Claude Code月租：20美元/月，抢钱呢', style='List Bullet')
doc.add_paragraph('GitHub Copilot：10美元/月，功能一般', style='List Bullet')
doc.add_paragraph('豆包：别用，浪费生命', style='List Bullet')

doc.add_heading('推荐用的（便宜够用）', 2)

p = doc.add_paragraph()
p.add_run('方案1：Cursor免费版 + DeepSeek').bold = True
doc.add_paragraph('Cursor有免费版，每月2000次补全够用了')
doc.add_paragraph('DeepSeek V4 Flash超便宜，1块钱能跑100万token')
doc.add_paragraph('我有glm4.7的周卡，需要的话找我拿')
doc.add_paragraph('不推荐方舟的codingplan，glm5和k2.5以上的模型够用')

p = doc.add_paragraph()
p.add_run('方案2：ccswitch切换模型（省钱技巧）').bold = True
doc.add_paragraph('平时用DeepSeek（便宜）')
doc.add_paragraph('遇到搞不定的bug，切到Claude（贵但强）')
doc.add_paragraph('这样能省80%的钱')

doc.add_page_break()

# 三、同学A干啥
doc.add_heading('三、同学A干啥（会点代码那位）', 1)

doc.add_heading('任务1：修3个bug（3天）', 2)
doc.add_paragraph('目标：让正确率从63%涨到75%')
doc.add_paragraph()
doc.add_paragraph('要改这3个文件：')
doc.add_paragraph('backend/agents/__init__.py - 让FeasibilityAgent更严格，该拒绝就拒绝', style='List Number')
doc.add_paragraph('backend/services/solver.py - 美食路线别跨区，限制10km内', style='List Number')
doc.add_paragraph('backend/services/filters.py - 凌晨2点别推荐咖啡店', style='List Number')

doc.add_heading('任务2：搭测试框架（2天）', 2)
doc.add_paragraph('写3个测试文件：')
doc.add_paragraph('test_layer1.py - 测单个函数对不对', style='List Bullet')
doc.add_paragraph('test_layer2.py - 测Agent之间配合是否正常', style='List Bullet')
doc.add_paragraph('test_layer3.py - 测完整流程，用LongCat免费评分', style='List Bullet')

doc.add_heading('任务3：批量生成测试场景（1天）', 2)
doc.add_paragraph('写个脚本自动生成100个场景：')
doc.add_paragraph('预算边界：0元、1元、99元、100元、101元...', style='List Bullet')
doc.add_paragraph('时间边界：23:59、00:00、00:01、凌晨3点...', style='List Bullet')
doc.add_paragraph('人工挑一遍，去掉太离谱的', style='List Bullet')

doc.add_page_break()

# 四、同学B干啥
doc.add_heading('四、同学B干啥（完全不会代码那位）', 1)

doc.add_heading('任务1：当"判官"（1天）', 2)
doc.add_paragraph('看11个失败的场景，判断是不是LLM在故意找茬')
doc.add_paragraph('比如系统给了合理路线，LLM非说"不够好"给低分，这种要标出来')
doc.add_paragraph('输出：X个真有问题，Y个是被冤枉的')

doc.add_heading('任务2：写真实场景（2天）', 2)
doc.add_paragraph('写20个你真的会用到的场景，别瞎编：')
doc.add_paragraph('例：周末带女朋友约会，预算300，想拍照好看')
doc.add_paragraph('格式：名称 + 你咋说的 + 期望啥结果 + 为啥真实')
doc.add_paragraph('覆盖不同人群：情侣、带娃、朋友聚会、老人')

doc.add_heading('任务3：当间谍（1天）', 2)
doc.add_paragraph('用同样的需求测竞品：')
doc.add_paragraph('美团、高德、小红书', style='List Bullet')
doc.add_paragraph('记录：他们推荐啥、好在哪、差在哪', style='List Bullet')
doc.add_paragraph('输出对比表，看看我们差在哪', style='List Bullet')

doc.add_page_break()

# 五、架构图咋维护
doc.add_heading('五、架构图咋维护', 1)
doc.add_paragraph('每次改完代码，架构图也要更新，不然过几天自己都忘了啥结构')
doc.add_paragraph()
doc.add_paragraph('自动化方案：')
doc.add_paragraph('每次git commit后，自动跑脚本生成新架构图', style='List Bullet')
doc.add_paragraph('存在docs/architecture.md（文字版）和docs/architecture.png（图片版）', style='List Bullet')
doc.add_paragraph('谁改了啥、哪里有问题，一目了然', style='List Bullet')

doc.add_page_break()

# 六、每周咋同步
doc.add_heading('六、每周咋同步', 1)

doc.add_heading('同学A汇报（看数据）', 2)
doc.add_paragraph('修了X个bug，正确率从63% → XX%')
doc.add_paragraph('API花了多少钱：￥X元（别超预算）')
doc.add_paragraph('下周准备干啥')

doc.add_heading('同学B汇报（看体验）', 2)
doc.add_paragraph('审了X个案例，Y个真有问题')
doc.add_paragraph('写了20个场景，覆盖了XX人群')
doc.add_paragraph('竞品对比发现：美团在XX方面比我们强')

doc.add_page_break()

# 七、工具清单
doc.add_heading('七、工具清单', 1)

table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '谁用'
hdr_cells[1].text = '用啥'
hdr_cells[2].text = '多少钱'

data = [
    ['同学A写代码', 'Cursor免费版 + DeepSeek', '免费 / 1元百万token'],
    ['同学A省钱', 'ccwitch切模型', '免费'],
    ['同学B测试', '不用编程工具', '-'],
    ['一起评分', 'LongCat（公测免费）', '免费']
]
for i, row_data in enumerate(data, 1):
    row_cells = table.rows[i].cells
    for j, text in enumerate(row_data):
        row_cells[j].text = text

doc.add_page_break()

# 八、别踩坑
doc.add_heading('八、别踩坑', 1)

doc.add_heading('花钱红线', 2)
doc.add_paragraph('别买Claude Code月租（20美元太贵）', style='List Bullet')
doc.add_paragraph('别买GitHub Copilot', style='List Bullet')
doc.add_paragraph('别在Cursor里开Pro（免费版够用）', style='List Bullet')
doc.add_paragraph('别用豆包', style='List Bullet')

doc.add_heading('协作红线', 2)
doc.add_paragraph('每天同步一次，别闷头干一周才说', style='List Bullet')
doc.add_paragraph('不懂就问，别猜', style='List Bullet')
doc.add_paragraph('改完代码必须跑测试，哪怕只跑Layer 1', style='List Bullet')
doc.add_paragraph('别直接改main分支，开feature分支', style='List Bullet')

doc.add_page_break()

# 九、第一周干啥
doc.add_heading('九、第一周干啥', 1)

doc.add_heading('Day 1（对齐）', 2)
doc.add_paragraph('[ ] 同学A：装Cursor，配置DeepSeek API', style='List Bullet')
doc.add_paragraph('[ ] 同学B：看失败案例分析文档', style='List Bullet')
doc.add_paragraph('[ ] 晚上：一起开会，确认分工', style='List Bullet')

doc.add_heading('Day 2-3（干活）', 2)
doc.add_paragraph('[ ] 同学A：改FeasibilityAgent，让它更严格', style='List Bullet')
doc.add_paragraph('[ ] 同学B：审11个失败案例，标出哪些被冤枉', style='List Bullet')
doc.add_paragraph('[ ] 晚上：同步进展', style='List Bullet')

doc.add_heading('Day 4-5（继续）', 2)
doc.add_paragraph('[ ] 同学A：写测试框架Layer 1和2', style='List Bullet')
doc.add_paragraph('[ ] 同学B：写20个真实场景', style='List Bullet')
doc.add_paragraph('[ ] 周五：跑一遍测试，看正确率涨没涨', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('就这些，干就完了。有问题随时群里@我。')

# 保存
output_path = 'C:\\Users\\wang\\Desktop\\CityFlow_干活手册.docx'
doc.save(output_path)
print(f'搞定：{output_path}')
