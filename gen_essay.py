"""生成架构演化随笔 docx。"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# ── 样式 ──
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.8
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# ═══════════════════════════════════════════
# 正文
# ═══════════════════════════════════════════

# 标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('从天轴到电动机')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('CityFlow C版本架构演化手记：一次从中央计划到分布式智能体的思想实验')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x88)
run.font.italic = True

doc.add_paragraph()  # 空行

# ── 引言 ──
doc.add_heading('一、缘起：一台跑不满的机器', level=1)

doc.add_paragraph(
    '2026年5月，珠海。CityFlow项目进入第三周。'
    'C版本——分布式智能体架构——刚完成第二轮优化，30场景LLM评测拿到5.5分，通过率18%。'
    '这个数字比A版本（3层联邦架构）的5.3分好了那么一点，但离及格线6.5分还差得远。'
)

doc.add_paragraph(
    '我盯着架构诊断图发呆。图上标注了三个红色问题：'
)

problems = [
    ('P0', 'coordinator把Agent精选的POI传给solver，solver的Phase 0用关键字过滤把它们全丢了。"安静画画"只匹配到1个POI，候选池从120个缩到2个，路线只有3-4站。'),
    ('P1', 'intent_match 5.5分，24%的场景得分≤3。"摄影夜景"选了白天景点，"亲子"选了剧本杀，"雨夜漫步"选了攀岩。'),
    ('P2', 'group_debate不能修改proposals，Annotated[list, operator.add]只允许追加。群聊标记的冲突靠name匹配删除，容易漏。'),
]
for label, desc in problems:
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
    p.add_run(desc)

doc.add_paragraph(
    '我试了很多东西：把Agent精选POI直传solver、清空scene_requirements避免过滤、'
    '补充地理邻近候选、甚至让solver完全接管选点。'
    '结果都是一样的——solver对C版本来说是负优化。'
    '纯Agent手动组装反而更好，带娃蹦迪场景从4分跳到6分。'
)

doc.add_paragraph(
    '这不是代码的问题。这是架构的问题。'
)

# ── 天轴 ──
doc.add_heading('二、天轴时代：一个中央蒸汽机驱动整座工厂', level=1)

doc.add_paragraph(
    '1880年代的工厂是什么样的？'
)

doc.add_paragraph(
    '一座巨大的蒸汽机坐落在厂房中央。'
    '从蒸汽机伸出一根主轴，横贯整个车间天花板。'
    '主轴上挂满皮带轮，每根皮带垂下来，连接到一台具体的机器——车床、铣床、织布机。'
)

doc.add_paragraph(
    '这根主轴，叫"天轴"（Line Shaft）。'
)

doc.add_paragraph(
    '天轴系统有几个致命问题：'
)

shaft_problems = [
    ('单一故障点', '主轴断了，全厂停工。一根皮带松了，整条线速度不稳。'),
    ('效率天花板', '所有机器必须同步运转。想单独加快一台车床？不行，主轴转速是固定的。想给角落的机器单独调速？需要加装一套复杂的锥轮变径装置。'),
    ('刚性耦合', '新增一台设备？工程师要重新计算整根天轴的载荷分布。设备的位置不能按工艺流程摆放，只能按"离天轴近"摆放。'),
    ('能量浪费', '离蒸汽机最远的机器，经过几十个皮带轮的摩擦损耗，实际得到的动力只剩60%。工程师知道问题在哪，但天轴的物理定律不允许改变。'),
]
for label, desc in shaft_problems:
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.font.bold = True
    p.add_run(desc)

doc.add_paragraph(
    '但当时没有别的选择。动力源只有一个——蒸汽机。'
    '所以工程师们把所有聪明才智都用在了"如何让天轴更高效"上：'
    '更精密的齿轮、更紧的皮带、更复杂的分轴系统。'
)

doc.add_paragraph(
    '像不像我们的solver？'
)

# ── Solver = 天轴 ──
doc.add_heading('三、Solver就是那根天轴', level=1)

doc.add_paragraph(
    '回过头看C版本的solver，它和天轴是同一个东西：'
)

doc.add_paragraph(
    'Solver是一台5阶段的路线优化"蒸汽机"。'
    '它从Phase 0开始选点——用关键字匹配过滤scene_requirements，'
    '然后Phase 1贪心初始化，Phase 2做2-opt，Phase 3插呼吸空间，Phase 4高潮收尾。'
    '整个管线精密、完整、自洽。'
)

doc.add_paragraph(
    '但问题是：7个Agent已经把活干了。'
)

doc.add_paragraph(
    'POI Agent调了DeepSeek LLM，根据用户是亲子还是情侣，选出了5-8个最合适的景点。'
    '它的poi_quality评分是6.6——整个系统里最高的。'
    'Food Agent选了1-3个餐厅，还标了午餐还是晚餐。'
    'Traffic Agent看了30个POI的位置，算了距离矩阵，排了个最优顺序。'
)

doc.add_paragraph(
    '然后这些精心挑选的结果被送到solver。'
    'Solver做了什么？它忽略Agent的选择，从原始候选池重新开始。'
    'Phase 0用"安静画画"四个字在4734个POI里找匹配——找到1个。'
    'Phase 1从这1个POI开始贪心——路线只有3站。'
    'Agent精选的12-18个提案，全部浪费。'
)

doc.add_paragraph(
    '这就是天轴的本质：不管每台"机器"（Agent）自己能干多好，'
    '它们都被迫接到同一根"主轴"（Solver）上，接受统一的转速（Phase 0-5的固定流程）。'
)

doc.add_paragraph(
    '我们尝试过修复天轴：把Agent精选的POI直接传给solver，'
    '清空scene_requirements跳过Phase 0，只让solver做Phase 2-5的排序。'
    '效果呢？路线长度从3-4站变成5-6站，但geo_continuity反而从5.0降到5.1。'
    '因为solver的时间窗约束太死——它认为一个POI只能待90分钟，'
    '超时就要惩罚，所以宁可少安排几个站也不让时间溢出。'
)

doc.add_paragraph(
    '修天轴是没有出路的。因为问题不在天轴本身，'
    '而在"只有一根天轴"这个架构假设。'
)

# ── 电动机革命 ──
doc.add_heading('四、1887年：每台机器获得自己的电动机', level=1)

doc.add_paragraph(
    '1887年，尼古拉·特斯拉发明了交流电动机。'
)

doc.add_paragraph(
    '突然之间，工厂不再需要天轴了。'
    '每台机器可以装上自己的电动机，按自己的速度运转，'
    '停机检修不影响其他机器，新增设备不需要重新计算载荷。'
)

doc.add_paragraph(
    '更重要的是：工厂布局改变了。'
    '以前设备必须沿天轴一字排开；现在设备可以按工艺流程自由摆放。'
    '车床可以紧挨着铣床，不是因为天轴正好经过这里，'
    '而是因为工序上"先车后铣"最合理。'
)

doc.add_paragraph(
    '电气化不是"更好的天轴"。它是一种全新的组织方式。'
    '它释放的不只是能量，是自由度。'
)

# ── Agent = 电动机 ──
doc.add_heading('五、每台Agent获得自己的LLM', level=1)

doc.add_paragraph(
    '我们的"电气化"发生在一个深夜。'
)

doc.add_paragraph(
    '那天我盯着Traffic Agent的prompt看。它只有三行：'
)

p = doc.add_paragraph()
run = p.add_run(
    '你是城市交通规划专家。根据景点分布给出最优交通方案和路线建议。\n'
    '1. 景点之间的距离和交通方式\n'
    '2. 推荐的游览顺序（考虑地理连贯性）\n'
    '3. 交通时间估算\n'
    'suggested_order按地理最优排列。'
)
run.font.name = 'Consolas'
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph(
    '这个prompt就像一根皮带——它把Traffic Agent接到"地理最优"这根天轴上。'
    'Agent不懂用户是亲子还是情侣，不懂上午精力好还是下午适合轻松，'
    '不懂"安静画画"意味着美术馆而不是蹦迪。'
    '它只懂一件事：最短路径。'
)

doc.add_paragraph(
    '我重写了Traffic Agent的prompt。这是"给它装上电动机"的瞬间：'
)

p = doc.add_paragraph()
run = p.add_run(
    '你是城市旅行路线规划专家。你需要设计一条高质量的一日游路线。\n\n'
    '你的核心任务：\n'
    '1. 【地理连贯】按区域聚类排序，避免来回折返\n'
    '2. 【时间节奏】遵循情绪曲线设计：\n'
    '   - 上午(9-12点)：精力充沛，安排主力景点\n'
    '   - 午餐(11:30-13:00)：选景点附近的特色餐饮\n'
    '   - 下午(13-17点)：次级景点或室内\n'
    '   - 傍晚(17-19点)：观景/休闲\n'
    '3. 【场景适配】亲子：景点间距要短 / 情侣：海滨浪漫路线\n'
    '4. 【高效交通】同区域景点连走'
)
run.font.name = 'Consolas'
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x22, 0x88, 0x44)

doc.add_paragraph(
    '现在Traffic Agent有了自己的"电动机"——它能理解用户意图、'
    '设计情绪曲线、考虑时间节奏、适配场景类型。'
    '它不再是一台只能算最短路径的机器，它是一个能做综合判断的规划师。'
)

doc.add_paragraph(
    '同样的改造发生在Food Agent：'
)

food_changes = [
    '加了坐标——每家餐厅有了经纬度，LLM能判断它离哪个景点近',
    '加了景点位置——LLM知道用户上午在横琴，不会推荐拱北的餐厅',
    '加了群体适配——亲子选有儿童餐的，情侣选氛围好的',
    '加了meal_time——午餐选上午景点附近的，晚餐选下午景点附近的',
]
for c in food_changes:
    doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph(
    'Hotel Agent加了坐标，Weather Agent去了硬编码的"5月"，'
    'Local Expert Agent加了去重（不要再推荐已选的热门景点），'
    'Insurance Agent加了住宿信息和POI类型风险推断。'
)

doc.add_paragraph(
    '7台Agent，7台独立的"电动机"。每台都能在自己的领域做深度决策，'
    '而不是被动地被一根天轴拖着转。'
)

# ── Coordinator的觉醒 ──
doc.add_heading('六、Coordinator的觉醒：从工头到指挥家', level=1)

doc.add_paragraph(
    '但是有一个问题。'
)

doc.add_paragraph(
    'Agent们都有了"电动机"，能独立做高质量决策了。'
    '可是coordinator——负责把所有Agent结果组装成最终路线的节点——'
    '还在用天轴时代的逻辑。'
)

doc.add_paragraph(
    '它做了什么？它拿到了Traffic Agent精心排好的顺序，'
    '然后调了一个叫_nearest_neighbor_sort_stops的函数，'
    '重新按最近邻排序了一遍。'
    'Traffic Agent考虑了情绪曲线、时间节奏、场景适配的排序——全部白费。'
)

doc.add_paragraph(
    '它还硬编码了餐饮插入位置：午餐固定插第2个POI后面，'
    '晚餐固定插第5个POI后面。不管那家餐厅实际在地图上的哪里。'
    '结果得月舫（吉大片区）被插到了横琴的海洋王国后面——跨了20公里去吃午饭。'
)

doc.add_paragraph(
    '这就像一座已经电气化的工厂，每台机器都有自己的电动机了，'
    '但工头还在用天轴时代的排班表——强制所有机器同步启停，'
    '强制每台设备按它在旧天轴上的位置运转。'
    '新的自由度被旧的调度方式完全浪费了。'
)

doc.add_paragraph(
    '用户的一句话点醒了我：'
)

p = doc.add_paragraph()
run = p.add_run('"排序这个阶段本身也应该使用Agent，而fallback仅仅只是备用方案。"')
run.font.bold = True
run.font.italic = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)

doc.add_paragraph(
    'coordinator不应该自己做排序。它应该是一个指挥家——'
    '决定信谁、怎么拼、什么顺序，然后把具体的编排交给一个专门的"编排Agent"。'
)

doc.add_paragraph(
    '于是我把coordinator改成这样：'
)

coordinator_steps = [
    '收集各Agent提案，应用群聊冲突决议',
    '按Agent类型分类：poi / food / hotel / traffic',
    '把所有提案送给一个LLM"编排Agent"，让它综合考虑地理、时间、情绪、餐饮位置来编排路线',
    'LLM失败时，才用规则引擎（_fallback_assemble）兜底',
    '最后调用文案生成服务',
]
for i, step in enumerate(coordinator_steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'Step {i}: ')
    run.font.bold = True
    p.add_run(step)

doc.add_paragraph(
    '编排Agent的prompt是整个系统里最关键的——'
    '它看到的不是原始数据，而是7个专业Agent的精选结果：'
    'POI Agent选好的景点（含坐标和评分），Food Agent选好的餐厅（含坐标和meal_time），'
    'Traffic Agent建议的顺序，距离矩阵。'
    '它的任务是综合这些信息，排出一条地理连贯、时间合理、餐饮就近的路线。'
)

doc.add_paragraph(
    '这和天轴时代的区别是什么？'
    '天轴时代，coordinator像工头一样强迫所有Agent服从自己的规则。'
    '现在，coordinator像指挥家一样，信任每个乐手（Agent）的专业能力，'
    '自己只负责协调节奏和配合。'
)

# ── 数据对比 ──
doc.add_heading('七、实测：电气化前后的工厂效率', level=1)

doc.add_paragraph(
    '5场景LLM评测，同一个评分模型（LongCat-Flash-Lite），'
    '同一套场景，同一个及格线（6.5分）：'
)

# 表格
table = doc.add_table(rows=8, cols=4)
table.style = 'Light Grid Accent 1'
headers = ['指标', '天轴时代\n(规则编排)', '电气化后\n(LLM编排)', '变化']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data = [
    ('通过率', '3/5', '4/5', '+1'),
    ('overall', '7.0', '7.2', '+0.2'),
    ('intent_match', '7.6', '8.0', '+0.4'),
    ('poi_quality', '7.4', '7.4', '持平'),
    ('geo_continuity', '6.2', '6.4', '+0.2'),
    ('scene_diversity', '7.8', '8.0', '+0.2'),
    ('差异化路线', '5/5', '5/5', '持平'),
]
for i, (metric, before, after, delta) in enumerate(data):
    table.rows[i+1].cells[0].text = metric
    table.rows[i+1].cells[1].text = before
    table.rows[i+1].cells[2].text = after
    table.rows[i+1].cells[3].text = delta

doc.add_paragraph()

doc.add_paragraph(
    '对比更早的30场景数据（solver时代），变化更显著：'
)

table2 = doc.add_table(rows=4, cols=4)
table2.style = 'Light Grid Accent 1'
headers2 = ['版本', '平均分', '通过率', '架构特征']
for i, h in enumerate(headers2):
    table2.rows[0].cells[i].text = h

data2 = [
    ('Solver全权负责', '5.5', '18% (5/28)', '天轴：Agent结果被solver Phase 0丢弃'),
    ('纯Agent+规则编排', '7.0', '60% (3/5)', '拆掉天轴，但coordinator还在用旧规则'),
    ('Agent+LLM编排', '7.2', '80% (4/5)', '电气化：每台Agent自主决策，coordinator只调度'),
]
for i, (ver, score, rate, note) in enumerate(data2):
    table2.rows[i+1].cells[0].text = ver
    table2.rows[i+1].cells[1].text = score
    table2.rows[i+1].cells[2].text = rate
    table2.rows[i+1].cells[3].text = note

doc.add_paragraph()

doc.add_paragraph(
    'intent_match从5.5到8.0的飞跃，本质上是"谁理解用户"的权力转移。'
    '天轴时代，是solver用关键字匹配理解用户——"安静画画"在POI数据库里搜不到就放弃。'
    '电气化后，是POI Agent用LLM理解用户——"安静画画"="美术馆/画室/安静的咖啡馆"。'
    '理解能力的下放，带来了匹配质量的飞跃。'
)

# ── 架构反思 ──
doc.add_heading('八、反思：为什么天轴的诱惑如此强大', level=1)

doc.add_paragraph(
    '回头看，天轴（solver）的诱惑力来自三个地方：'
)

doc.add_heading('1. 优雅的理论', level=3)

doc.add_paragraph(
    'Solver有一个完美的5阶段流程：选点→贪心初始化→2-opt优化→呼吸空间→高潮收尾。'
    '每一步都有数学基础，TSP是运筹学的经典问题。'
    '这个优雅的理论让人很难承认它在实际场景中可能是负优化。'
)

doc.add_heading('2. 确定性幻觉', level=3)

doc.add_paragraph(
    'Solver是确定性的——同样的输入永远给同样的输出。'
    '这让人觉得它比LLM"可靠"。但实际上，'
    '确定性不等于正确性。'
    'solver确定性地把"安静画画"匹配到0个POI，这个结果是确定的，但也是灾难性的。'
    'LLM可能偶尔把"安静画画"理解为"安静的咖啡馆"而不是"美术馆"，'
    '但这个"不确定性"比solver的"确定性错误"好得多。'
)

doc.add_heading('3. 中央控制的安全感', level=3)

doc.add_paragraph(
    '一个中央的solver让人感到安心——一切尽在掌握。'
    '7个Agent各自为政的感觉像是"失控"。'
    '但"中央控制"和"中央瓶颈"之间只有一线之隔。'
    '当中央节点的处理能力无法匹配合法的复杂度时，'
    '中央控制就从优势变成了障碍。'
)

doc.add_paragraph(
    '这三个诱惑，和1880年代工厂主不愿放弃天轴的理由一模一样：'
    '"天轴系统我已经投入了大量资金（工程时间），"'
    '"我知道它的效率上限在哪里（可预测），"'
    '"那些电动机每台都有自己的脾气（不确定性）。"'
)

# ── 未完 ──
doc.add_heading('九、未完的工程', level=1)

doc.add_paragraph(
    '电气化不是一蹴而就的。1887年特斯拉发明交流电动机后，'
    '美国的工厂又花了20年才完成全面电气化。'
    '很多工厂甚至经历了"部分电气化"——'
    '保留天轴的同时给几台关键设备装上电动机，结果两套系统互相干扰，效率更差。'
)

doc.add_paragraph(
    '我们的C版本也还在这个过渡期。geo_continuity从6.2到6.4，提升不大。'
    '原因在于POI Agent选择的景点本身有重叠——珠海渔女、得月舫在5个场景里都出现。'
    '这不是编排能解决的问题，而是POI Agent的"电动机"还不够强。'
    '它需要更深层的场景特化——亲子不选渔女（太网红太挤），'
    '美食不选渔女（不是景点），特种兵选圆明新园（知名地标但经常被漏掉）。'
)

doc.add_paragraph(
    '下一步的工作是：'
)

next_steps = [
    'POI Agent的"电动机"继续强化——更多场景分化prompt，减少POI重叠',
    '编排Agent的prompt继续打磨——目前geo_continuity 6.4说明地理优化还有空间',
    '群聊协议增强——让group_debate不只是检测冲突，还要能做创造性优化（如"上午A下午B互补"）',
    '跑30场景全面评测——5场景的样本太小，需要更大规模验证',
]
for step in next_steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_paragraph(
    '但方向已经清楚了：'
)

p = doc.add_paragraph()
run = p.add_run(
    '不要再修天轴了。给每台Agent装上足够强的电动机，然后相信它们。'
)
run.font.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# ── 尾注 ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('2026年5月13日，珠海')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('CityFlow项目 · C版本架构演化记录')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.italic = True

# 保存
output_path = r'C:\Users\wang\Desktop\从天轴到电动机_CityFlow架构演化手记.docx'
doc.save(output_path)
print(f'已保存: {output_path}')
